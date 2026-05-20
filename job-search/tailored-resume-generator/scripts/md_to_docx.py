#!/usr/bin/env python3
"""Convert a tailored resume markdown file to a formatted DOCX.

Handles `**bold**` markdown syntax by converting it to native Word bold runs.

Usage:
    source /tmp/.venv/bin/activate  (if not already active)
    python3 scripts/md_to_docx.py <input.md> <output.docx>

Prerequisites (in /tmp/.venv):
    uv pip install python-docx lxml

If /tmp/.venv doesn't exist:
    cd /tmp && uv venv .venv && source .venv/bin/activate && uv pip install python-docx lxml

Markdown conventions expected by this script:
  - Name: ALL CAPS, no special chars, first line
  - Contact: line with @ symbol (centered, gray)
  - Section headers: ALL CAPS, > 5 chars, standalone
  - Job entries: **Bold Title** | Company | Location | Date
  - Skill lines: **Category:** items (plain paragraph, not bullet)
  - Achievement bullets: - text starting with dash (List Bullet style)
  - **text** in any line → native bold formatting
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


DARK  = RGBColor(0x1A, 0x1A, 0x1A)  # section headers
BODY  = RGBColor(0x33, 0x33, 0x33)  # normal text
MUTED = RGBColor(0x55, 0x55, 0x55)  # secondary (company, location)
FINE  = RGBColor(0x66, 0x66, 0x66)  # date


def _add_run(paragraph, text, bold=False, size=10, color=BODY):
    """Add a single run with formatting to a paragraph."""
    r = paragraph.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def _add_rich(paragraph, text, size=10, color=BODY):
    """Add text, converting **bold** markers to native bold runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _add_run(paragraph, part[2:-2], bold=True, size=size, color=color)
        else:
            _add_run(paragraph, part, bold=False, size=size, color=color)


def convert(input_path: str, output_path: str) -> None:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY

    with open(input_path) as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        raw = lines[i].strip()

        # Strip markdown header prefixes (# ## ### etc.)
        if raw.startswith('#'):
            raw = raw.lstrip('#').strip()
        if not raw:
            i += 1
            continue

        # ── Name line ──────────────────────────────────────────────
        if raw.isupper() and len(raw) < 50 and ':' not in raw and '|' not in raw:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, raw, bold=True, size=16)
            i += 1
            continue

        # ── Contact line ───────────────────────────────────────────
        if '@' in raw:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, raw, size=9.5, color=MUTED)
            i += 1
            continue

        # ── Section headers ────────────────────────────────────────
        if raw.isupper() and len(raw) > 5 and not raw.startswith('-') and '|' not in raw:
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            p.space_after = Pt(2)
            _add_run(p, raw, bold=True, size=11, color=DARK)
            i += 1
            continue

        # ── Skill category lines: "- **Category:** items" ─────────
        # Plain paragraph (not bullet) with bold category name.
        if raw.startswith('- ') and '**' in raw and ':' in raw[:60]:
            text = raw.lstrip('-• ').strip()
            p = doc.add_paragraph()
            p.space_before = Pt(2)
            p.space_after = Pt(1)
            # Split on first colon
            colon_idx = text.index(':')
            cat = text[:colon_idx].strip()
            rest = text[colon_idx + 1:]
            # Strip leftover ** markers
            cat = cat.replace('**', '')
            rest = rest.replace('**', '').strip()
            _add_run(p, cat, bold=True, size=10)
            _add_run(p, ': ' + rest, size=10)
            i += 1
            continue

        # ── Achievement bullets ────────────────────────────────────
        if raw.startswith('- ') or raw.startswith('• '):
            text = raw.lstrip('-• ').strip()
            p = doc.add_paragraph(style='List Bullet')
            p.space_before = Pt(0)
            p.space_after = Pt(1)
            _add_rich(p, text, size=10)
            i += 1
            continue

        # ── Job entry: "**Title** | Company | Location | Date" ────
        if '|' in raw:
            parts = [x.strip() for x in raw.split('|')]
            parts = [p for p in parts if p]
            p = doc.add_paragraph()
            p.space_before = Pt(5)
            p.space_after = Pt(0)
            for idx, part in enumerate(parts):
                if idx > 0:
                    _add_run(p, '  |  ', size=10.5, color=MUTED)
                _add_rich(p, part, size=10.5)
            # Consume any following lines that are not bullets/section headers
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt or nxt.startswith('-') or nxt.isupper() or '|' in nxt:
                    break
                j += 1
            i = j
            continue

        # ── Plain text (summary, cert line, etc.) ─────────────────
        p = doc.add_paragraph(raw)
        p.space_before = Pt(1)
        p.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(10)

        i += 1

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
